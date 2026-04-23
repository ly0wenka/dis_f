import argparse
import datetime as _dt
import re
from dataclasses import dataclass
from typing import Iterable, List, Optional, Tuple

import math2docx
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def _strip_comments(tex: str) -> str:
    out_lines: List[str] = []
    for line in tex.splitlines():
        buf = []
        escaped = False
        for ch in line:
            if ch == "%" and not escaped:
                break
            if ch == "\\" and not escaped:
                escaped = True
                buf.append(ch)
                continue
            escaped = False
            buf.append(ch)
        out_lines.append("".join(buf))
    return "\n".join(out_lines)


def _find_balanced_braces(s: str, open_brace_index: int) -> Tuple[str, int]:
    if open_brace_index >= len(s) or s[open_brace_index] != "{":
        raise ValueError("Expected '{' at index")
    depth = 0
    i = open_brace_index
    content_start = open_brace_index + 1
    while i < len(s):
        ch = s[i]
        if ch == "{":
            depth += 1
        elif ch == "}":
            depth -= 1
            if depth == 0:
                return s[content_start:i], i + 1
        i += 1
    raise ValueError("Unbalanced braces")


def _extract_command_arg(tex: str, command: str) -> Optional[str]:
    m = re.search(rf"\\{re.escape(command)}\s*\{{", tex)
    if not m:
        return None
    brace_idx = m.end() - 1
    content, _ = _find_balanced_braces(tex, brace_idx)
    return content.strip()


def _normalize_title_field(s: str) -> str:
    s = s.replace("\\\\", "\n")
    s = re.sub(r"\\large\s*\{([^}]*)\}", r"\1", s)
    lines = [re.sub(r"\s+", " ", ln).strip() for ln in s.splitlines()]
    s = "\n".join([ln for ln in lines if ln != ""]).strip()
    return s


UA_MONTHS = {
    1: "січня",
    2: "лютого",
    3: "березня",
    4: "квітня",
    5: "травня",
    6: "червня",
    7: "липня",
    8: "серпня",
    9: "вересня",
    10: "жовтня",
    11: "листопада",
    12: "грудня",
}


def _format_ua_date(d: _dt.date) -> str:
    return f"{d.day} {UA_MONTHS[d.month]} {d.year} р."


def _split_doc_body(tex: str) -> str:
    m1 = re.search(r"\\begin\s*\{\s*document\s*\}", tex)
    m2 = re.search(r"\\end\s*\{\s*document\s*\}", tex)
    if not m1 or not m2 or m2.start() <= m1.end():
        raise ValueError("Could not find document body (begin/end document).")
    return tex[m1.end() : m2.start()]


def _extract_tag(math: str) -> Tuple[str, Optional[str]]:
    tags = re.findall(r"\\tag\s*\{([^}]*)\}", math)
    math_clean = re.sub(r"\\tag\s*\{[^}]*\}", "", math)
    tag = tags[-1].strip() if tags else None
    return math_clean.strip(), tag


def _unwrap_boxed(math: str) -> str:
    # Replace occurrences of \boxed{...} with ...
    out = []
    i = 0
    while i < len(math):
        if math.startswith("\\boxed", i):
            j = i + len("\\boxed")
            while j < len(math) and math[j].isspace():
                j += 1
            if j < len(math) and math[j] == "{":
                inner, next_i = _find_balanced_braces(math, j)
                out.append(inner)
                i = next_i
                continue
        out.append(math[i])
        i += 1
    return "".join(out)


def _sanitize_math(math: str, *, drop_alignment_markers: bool = False) -> str:
    math = math.strip()
    math, _ = _extract_tag(math)
    math = _unwrap_boxed(math)
    if drop_alignment_markers:
        math = math.replace("&", "")
    # latex2mathml doesn't always like \text{...}; map to \mathrm{...}
    math = re.sub(r"\\text\s*\{([^}]*)\}", r"\\mathrm{\1}", math)
    # Common spacing commands -> plain space
    math = re.sub(r"\\[,;!]\s*", " ", math)
    math = re.sub(r"\\quad\b", " ", math)
    math = re.sub(r"\\qquad\b", " ", math)
    math = re.sub(r"\s+", " ", math).strip()
    return math


def _add_math_paragraph(doc: Document, math: str, tag: Optional[str], *, drop_alignment_markers: bool = False) -> None:
    p = doc.add_paragraph()
    try:
        math2docx.add_math(p, _sanitize_math(math, drop_alignment_markers=drop_alignment_markers))
    except Exception:
        # Fallback: write LaTeX source as plain text so nothing is lost.
        p.add_run(math)
        if tag:
            p.add_run(f" ({tag})")
        return
    if tag:
        p.add_run(f" ({tag})")


def _add_page_break(doc: Document) -> None:
    doc.add_page_break()


def _parse_inline_runs(text: str) -> List[Tuple[str, str]]:
    """
    Returns list of (kind, value):
      kind in {"text", "math", "bold_text"}.
    Supports:
      - inline math: $...$ and \\(...\\)
      - bold text: \\textbf{...}
    """
    runs: List[Tuple[str, str]] = []
    i = 0

    def emit(kind: str, val: str) -> None:
        if not val:
            return
        runs.append((kind, val))

    while i < len(text):
        if text.startswith("\\textbf", i):
            j = i + len("\\textbf")
            while j < len(text) and text[j].isspace():
                j += 1
            if j < len(text) and text[j] == "{":
                inner, next_i = _find_balanced_braces(text, j)
                emit("bold_text", inner)
                i = next_i
                continue
        if text.startswith("\\(", i):
            j = text.find("\\)", i + 2)
            if j != -1:
                emit("math", text[i + 2 : j])
                i = j + 2
                continue
        if text[i] == "$":
            # handle $$...$$ as block-ish inline; still emit math
            if i + 1 < len(text) and text[i + 1] == "$":
                j = text.find("$$", i + 2)
                if j != -1:
                    emit("math", text[i + 2 : j])
                    i = j + 2
                    continue
            j = i + 1
            while j < len(text):
                if text[j] == "$" and text[j - 1] != "\\":
                    emit("math", text[i + 1 : j])
                    i = j + 1
                    break
                j += 1
            else:
                emit("text", text[i:])
                i = len(text)
            continue
        # plain text until next special token
        next_pos = len(text)
        for token in ["\\textbf", "\\(", "$"]:
            p2 = text.find(token, i + 1)
            if p2 != -1:
                next_pos = min(next_pos, p2)
        emit("text", text[i:next_pos])
        i = next_pos
    return runs


def _add_rich_paragraph(doc: Document, text: str, style: Optional[str] = None) -> None:
    text = text.strip()
    if not text:
        return
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    p.paragraph_format.space_after = 0
    for kind, val in _parse_inline_runs(text):
        if kind == "text":
            p.add_run(val)
        elif kind == "bold_text":
            r = p.add_run(val)
            r.bold = True
        elif kind == "math":
            try:
                math2docx.add_math(p, _sanitize_math(val))
            except Exception:
                p.add_run(val)


def _split_table_cells(row: str) -> List[str]:
    """
    Split a LaTeX table row into cells, avoiding splits inside math mode.
    """
    cells: List[str] = []
    buf: List[str] = []

    i = 0
    in_math = False
    in_display_math = False
    in_paren_math = False

    def flush() -> None:
        val = "".join(buf).strip()
        buf.clear()
        cells.append(val)

    while i < len(row):
        if row.startswith("\\(", i):
            in_paren_math = True
            buf.append("\\(")
            i += 2
            continue
        if row.startswith("\\)", i):
            in_paren_math = False
            buf.append("\\)")
            i += 2
            continue
        if row.startswith("$$", i) and not in_paren_math:
            in_display_math = not in_display_math
            buf.append("$$")
            i += 2
            continue
        ch = row[i]
        if ch == "$" and not in_paren_math:
            # toggle inline math
            in_math = not in_math
            buf.append(ch)
            i += 1
            continue
        if ch == "&" and not in_math and not in_display_math and not in_paren_math:
            flush()
            i += 1
            continue
        buf.append(ch)
        i += 1
    flush()
    return cells


def _emit_longtable_as_paragraphs(doc: Document, lines: List[str], start_index: int) -> int:
    """
    Convert \\begin{longtable}...\\end{longtable} into plain paragraphs (no Word tables).
    Returns the next line index after \\end{longtable}.
    """
    i = start_index
    if not re.match(r"\\begin\{longtable\}\s*\{", lines[i].strip()):
        raise ValueError("Expected longtable begin at start_index.")

    caption: Optional[str] = None
    i += 1

    row_buf: List[str] = []

    def finalize_row(row_text: str) -> None:
        row_text = row_text.strip()
        if not row_text:
            return
        row_text = re.sub(r"\\\\\s*$", "", row_text).strip()
        row_text = row_text.replace("\\hfill", " ")
        row_text = re.sub(r"\s+", " ", row_text).strip()
        if not row_text:
            return

        cells = _split_table_cells(row_text)
        cells = [c.strip() for c in cells if c.strip()]
        if not cells:
            return

        # Heuristics: render common 2/3-column tables as readable text lines.
        if len(cells) == 2:
            _add_rich_paragraph(doc, f"{cells[0]} — {cells[1]}", style="List Bullet")
        elif len(cells) == 3:
            lead = cells[0]
            if re.fullmatch(r"\d+", lead):
                _add_rich_paragraph(doc, f"{cells[1]} — {cells[2]}", style="List Number")
            else:
                _add_rich_paragraph(doc, f"{cells[0]} | {cells[1]} | {cells[2]}")
        else:
            _add_rich_paragraph(doc, " | ".join(cells))

    while i < len(lines):
        s = lines[i].strip()
        if s.startswith("\\end{longtable}"):
            break

        if not s or s == "\\hline":
            i += 1
            continue

        if s.startswith("\\caption"):
            try:
                cap, _ = _find_balanced_braces(s, s.find("{"))
                caption = cap.strip()
            except Exception:
                caption = None
            i += 1
            continue

        if s.startswith("\\label") or s.startswith("\\addcontentsline"):
            i += 1
            continue

        # Accumulate row lines until the LaTeX row terminator \\ is reached.
        row_buf.append(s)
        combined = " ".join(row_buf).strip()
        if re.search(r"\\\\\s*$", combined):
            finalize_row(combined)
            row_buf.clear()
        i += 1

    # Flush any unterminated row
    if row_buf:
        finalize_row(" ".join(row_buf))
        row_buf.clear()

    if caption:
        p = doc.add_paragraph()
        r = p.add_run(caption)
        r.italic = True

    # Skip the \end{longtable} line
    while i < len(lines) and "\\end{longtable}" not in lines[i]:
        i += 1
    if i < len(lines):
        i += 1
    return i


def _set_default_font(doc: Document, font_name: str = "Times New Roman", font_size_pt: int = 12) -> None:
    style = doc.styles["Normal"]
    font = style.font
    font.name = font_name
    font.size = Pt(font_size_pt)


@dataclass
class Preamble:
    title: Optional[str]
    author: Optional[str]
    date_raw: Optional[str]


def _parse_preamble(tex: str) -> Preamble:
    preamble = tex.split("\\begin{document}", 1)[0]
    title = _extract_command_arg(preamble, "title")
    author = _extract_command_arg(preamble, "author")
    date_raw = _extract_command_arg(preamble, "date")
    return Preamble(
        title=_normalize_title_field(title) if title else None,
        author=author,
        date_raw=date_raw,
    )


def _emit_title_block(doc: Document, preamble: Preamble) -> None:
    if preamble.title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(preamble.title)
        r.bold = True
    if preamble.author:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(preamble.author)
    if preamble.date_raw:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if preamble.date_raw.strip() == "\\today":
            p.add_run(_format_ua_date(_dt.date.today()))
        else:
            p.add_run(preamble.date_raw.strip())


def convert_tex_to_docx(input_tex: str, output_docx: str) -> None:
    raw = open(input_tex, "r", encoding="utf-8").read()
    raw = _strip_comments(raw)
    preamble = _parse_preamble(raw)
    body = _split_doc_body(raw)

    doc = Document()
    _set_default_font(doc)

    lines = body.splitlines()
    i = 0
    para_buf: List[str] = []

    def flush_paragraph() -> None:
        nonlocal para_buf
        txt = " ".join(x.strip() for x in para_buf if x.strip())
        para_buf = []
        if txt:
            _add_rich_paragraph(doc, txt)

    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if not stripped:
            flush_paragraph()
            i += 1
            continue

        if stripped == "\\maketitle":
            flush_paragraph()
            _emit_title_block(doc, preamble)
            i += 1
            continue

        if stripped.startswith("\\tableofcontents"):
            flush_paragraph()
            i += 1
            continue

        if stripped.startswith("\\newpage"):
            flush_paragraph()
            _add_page_break(doc)
            i += 1
            continue

        m = re.match(r"\\(subsubsection|subsection|section)\*?\s*\{", stripped)
        if m:
            flush_paragraph()
            content, _ = _find_balanced_braces(stripped, stripped.find("{"))
            level = {"section": 1, "subsection": 2, "subsubsection": 3}[m.group(1)]
            doc.add_heading(content.strip(), level=level)
            i += 1
            continue

        if stripped.startswith("\\begin{abstract}"):
            flush_paragraph()
            abstract_lines: List[str] = []
            i += 1
            while i < len(lines) and "\\end{abstract}" not in lines[i]:
                abstract_lines.append(lines[i])
                i += 1
            doc.add_heading("Анотація", level=2)
            _add_rich_paragraph(doc, " ".join(x.strip() for x in abstract_lines if x.strip()))
            while i < len(lines) and "\\end{abstract}" not in lines[i]:
                i += 1
            i += 1
            continue

        env_m = re.match(r"\\begin\{(equation\*?|align\*?)\}", stripped)
        if env_m:
            flush_paragraph()
            env = env_m.group(1)
            i += 1
            math_lines: List[str] = []
            end_pat = rf"\\end\{{{re.escape(env)}\}}"
            while i < len(lines) and not re.search(end_pat, lines[i]):
                math_lines.append(lines[i])
                i += 1
            i += 1  # skip \end

            math_content = "\n".join(math_lines).strip()
            if env.startswith("align"):
                parts = [p.strip() for p in re.split(r"\\\\\s*", math_content) if p.strip()]
                for part in parts:
                    part_clean, tag = _extract_tag(part)
                    _add_math_paragraph(doc, part_clean, tag, drop_alignment_markers=True)
            else:
                math_clean, tag = _extract_tag(math_content)
                _add_math_paragraph(doc, math_clean, tag)
            continue

        if re.match(r"\\begin\{longtable\}\s*\{", stripped):
            flush_paragraph()
            i = _emit_longtable_as_paragraphs(doc, lines, i)
            continue

        if stripped.startswith("\\addcontentsline"):
            flush_paragraph()
            i += 1
            continue

        if stripped.startswith("\\begin{itemize}") or stripped.startswith("\\begin{enumerate}"):
            flush_paragraph()
            is_enum = stripped.startswith("\\begin{enumerate}")
            end = "\\end{enumerate}" if is_enum else "\\end{itemize}"
            style = "List Number" if is_enum else "List Bullet"
            i += 1
            while i < len(lines) and end not in lines[i]:
                item_line = lines[i].strip()
                if item_line.startswith("\\item"):
                    item_text = item_line[len("\\item") :].strip()
                    _add_rich_paragraph(doc, item_text, style=style)
                i += 1
            i += 1
            continue

        # Default: treat as text content.
        para_buf.append(stripped)
        i += 1

    flush_paragraph()
    doc.save(output_docx)


def main() -> None:
    ap = argparse.ArgumentParser(description="Convert LaTeX .tex to DOCX with native Word equations via math2docx.")
    ap.add_argument("input_tex", help="Path to input .tex file")
    ap.add_argument(
        "-o",
        "--output",
        dest="output_docx",
        help="Path to output .docx file",
        default=None,
    )
    args = ap.parse_args()

    output = args.output_docx
    if output is None:
        if args.input_tex.lower().endswith(".tex"):
            output = args.input_tex[:-4] + "_math2docx.docx"
        else:
            output = args.input_tex + "_math2docx.docx"

    convert_tex_to_docx(args.input_tex, output)
    print(output)


if __name__ == "__main__":
    main()
